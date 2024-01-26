import psycopg2
from psycopg2 import sql
from psycopg2.extras import DictCursor
from docx import Document
import pandas as pd
import os

class ReadDatabase:
    def __init__(self):
        self.input_file_path = ""
        self.out_file_path = ""
        self.host = ""
        self.database = ""
        self.user = ""
        self.password = ""
        self.port = ""

# 建立连接
    def connection(self):
        db_params = {
            'host': self.host,
            'database': self.database,
            'user': self.user,
            'password': self.password,
            'port': self.port
        }

        # Establish a connection to the PostgreSQL database
        connection = psycopg2.connect(**db_params)
        return connection

# 获取数据库表名
    def get_ds_table_names(self):
        connection = self.connection()
        # Create a cursor object to execute SQL queries with DictCursor for easy access to column names
        cursor = connection.cursor(cursor_factory=DictCursor)

        # Get the list of tables in the database
        table_list_query = sql.SQL("SELECT table_name FROM information_schema.tables WHERE table_schema = 'public'")
        cursor.execute(table_list_query)
        tables = [table[0] for table in cursor.fetchall()]
        df = pd.DataFrame(tables, columns=["表名"])
        df.to_excel(self.out_file_path, index=False)

# 读取表名列表，获取表格元数据写入docx文件中
    def get_tables_meta_data(self):
        tb_list = pd.read_excel(self.input_file_path)["表名"].to_list()
        tb_name_list = pd.read_excel(self.input_file_path)["别名"].to_list()
        connection = self.connection()
        cursor = connection.cursor(cursor_factory=DictCursor)
        # Create a Document object for the DOCX file
        doc = Document()

        # Iterate through each table and retrieve metadata
        for table_name, alias_name in zip(tb_list, tb_name_list):
            # Retrieve metadata for the current table
            # 获取字段名称、备注、是否主外键
            metadata_query = sql.SQL(f"""
                SELECT
                    a.attname AS column_name,
                    d.description AS column_remark,
                    format_type(a.atttypid, a.atttypmod) AS data_type,
                    CASE
                        WHEN con.contype = 'p' THEN '主'
                        WHEN con.contype = 'f' THEN '外'
                        ELSE ''
                    END AS key_type
                FROM
                    pg_attribute a
                LEFT JOIN
                    pg_description d ON a.attrelid = d.objoid AND a.attnum = d.objsubid
                LEFT JOIN
                    pg_constraint con ON con.conrelid = a.attrelid AND a.attnum = ANY(con.conkey)
                WHERE
                    a.attrelid = '"public"."{table_name}"'::regclass
                    AND a.attnum > 0
            """)
            try:
                # 获取表格字段元数据
                cursor.execute(metadata_query)
                matadata_values = cursor.fetchall()
                #print(table_name)
                #print(matadata_values)
                # 列表储存新的表格字段元数据
                metadata_values_list = []
                for item in matadata_values:
                    col_name = item[0]
                    col_dp = item[2]

                    # 判断字段是否为空
                    isnull_query = sql.SQL(f"""
                        SELECT {col_name}
                        FROM {table_name};
                    """)
                    cursor.execute(isnull_query)
                    isnull = cursor.fetchone()[0]

                    if isnull is not None:
                        col_is_null = "√"
                    else:
                        col_is_null = ""
                    #print(col_is_null)
                    # 获取字符字段的值域
                    if col_name not in('id', "ocean_data_identification_code", "administrative_division",
                                       "administrative_division_code", "longitude", "latitude") and \
                            col_dp.startswith("character varying"):
                        # 获取字段唯一值
                        distinct_value_query = sql.SQL(f"""
                            SELECT DISTINCT {col_name}
                            FROM {table_name}
                        """)

                        cursor.execute(distinct_value_query)
                        distinct_values = cursor.fetchall()
                        dis_value_list = []
                        for val in distinct_values:
                            if val[0] is not None:
                                dis_value_list.append(val[0])
                        if len(dis_value_list) > 20:
                            dis_value_list = dis_value_list[:20]
                        distinct_values = "、".join(dis_value_list)
                    else:
                        distinct_values = ""
                    #print(distinct_values)
                    item_list = list(item)
                    item_list.append(col_is_null)
                    item_list.append(distinct_values)
                    #print(item_list)
                    metadata_values_list.append(item_list)
                #print(metadata_values_list)

                # Add table name as heading in the document
                doc.add_heading(f"Table: {alias_name}", level=1)

                # Create a table in the document for metadata
                table = doc.add_table(rows=1, cols=6)
                table.autofit = False

                # Add column headers to the table
                column_headers = table.rows[0].cells
                column_headers[0].text = "字段名"
                column_headers[1].text = "别名"
                column_headers[2].text = "数据类型"
                column_headers[3].text = "是否主外键"
                column_headers[4].text = "是否为空"
                column_headers[5].text = "值域/默认值"

                # Add metadata for each column to the table
                for column_info in metadata_values_list:
                    row_cells = table.add_row().cells
                    row_cells[0].text = column_info[0]
                    if column_info[1] is None:
                        row_cells[1].text = ""
                    else:
                        row_cells[1].text = column_info[1]

                    if column_info[2] is None:
                        row_cells[2].text = ""
                    else:
                        row_cells[2].text = column_info[2]
                    row_cells[3].text = str(column_info[3])
                    row_cells[4].text = column_info[4]

                    if column_info[0] in ["smuserid", "smarea", "smperimeter"]:
                        row_cells[5].text = "0"
                    elif column_info[5] is None:
                        row_cells[5].text = ""
                    else:
                        row_cells[5].text = column_info[5]

            except Exception as e:
                print(str(e))
                connection = self.connection()
                cursor = connection.cursor(cursor_factory=DictCursor)
                continue

        # Save the DOCX file
        doc.save(self.out_file_path)

        # Close the cursor and connection
        cursor.close()
        connection.close()

    # 读取表名列表，获取表格元数据写入excel文件中
    def get_tables_meta_data_excel(self):
        tb_list = pd.read_excel(self.input_file_path)["表名"].to_list()
        tb_name_list = pd.read_excel(self.input_file_path)["别名"].to_list()
        connection = self.connection()
        cursor = connection.cursor(cursor_factory=DictCursor)


        # Iterate through each table and retrieve metadata
        for table_name, alias_name in zip(tb_list, tb_name_list):
            # Retrieve metadata for the current table
            # 获取字段名称、备注、是否主外键
            metadata_query = sql.SQL(f"""
                SELECT
                    a.attname AS column_name,
                    d.description AS column_remark,
                    format_type(a.atttypid, a.atttypmod) AS data_type,
                    CASE
                        WHEN con.contype = 'p' THEN '主'
                        WHEN con.contype = 'f' THEN '外'
                        ELSE NULL
                    END AS key_type
                FROM
                    pg_attribute a
                LEFT JOIN
                    pg_description d ON a.attrelid = d.objoid AND a.attnum = d.objsubid
                LEFT JOIN
                    pg_constraint con ON con.conrelid = a.attrelid AND a.attnum = ANY(con.conkey)
                WHERE
                    a.attrelid = '"public"."{table_name}"'::regclass
                    AND a.attnum > 0
            """)
            try:
                # 获取表格字段元数据
                cursor.execute(metadata_query)
                matadata_values = cursor.fetchall()
                print(table_name)
                #print(matadata_values)
                # 列表储存新的表格字段元数据
                metadata_values_list = []
                for item in matadata_values:
                    col_name = item[0]
                    col_dp = item[2]

                    # 判断字段是否为空
                    isnull_query = sql.SQL(f"""
                        SELECT COUNT ({col_name}) = 0
                        FROM "{table_name}";
                    """)
                    cursor.execute(isnull_query)
                    isnull = cursor.fetchone()[0]

                    if isnull == False:
                        col_is_null = "√"
                    else:
                        col_is_null = ""
                    # print(col_is_null)
                    # 获取字符字段的值域
                    if col_name not in ('id', "ocean_data_identification_code", "administrative_division",
                                        "administrative_division_code", "longitude", "latitude") and \
                            col_dp.startswith("character varying"):
                        # 获取字段唯一值
                        distinct_value_query = sql.SQL(f"""
                            SELECT DISTINCT {col_name}
                            FROM "{table_name}"
                        """)

                        cursor.execute(distinct_value_query)
                        distinct_values = cursor.fetchall()
                        dis_value_list = []
                        for val in distinct_values:
                            if val[0] is not None:
                                dis_value_list.append(val[0])
                        dis_num = len(dis_value_list)
                        if dis_num > 20:
                            dis_value_list = dis_value_list[:20]
                        distinct_values = "、".join(dis_value_list)

                    else:
                        distinct_values = ""
                    # print(distinct_values)
                    item_list = list(item)
                    item_list.append(col_is_null)
                    item_list.append(distinct_values)
                    #print(item_list)
                    metadata_values_list.append(item_list)

                df = pd.DataFrame(metadata_values_list, columns=["字段", "备注", "数据类型", "主外键", "非空", "值域"])
                df.to_excel(os.path.join(self.out_file_path, alias_name+".xlsx"), index=False)
            except Exception as e:
                print(str(e))
                connection = self.connection()
                cursor = connection.cursor(cursor_factory=DictCursor)
                continue

        # Close the cursor and connection
        cursor.close()
        connection.close()


# 获取表格行数和大小
    def get_tables_size(self):
        tb_list = pd.read_excel(self.input_file_path)["表名"].to_list()
        connection = self.connection()
        cursor = connection.cursor(cursor_factory=DictCursor)

        row_count_list = []
        size_list = []
        for tb in tb_list:
            # 获取表格行数
            cursor.execute(f"SELECT COUNT(*) FROM {tb};")
            row_count = cursor.fetchone()[0]
            row_count_list.append(row_count)

            # 获取表格大小，以MB为单位
            cursor.execute(f"SELECT pg_total_relation_size('{tb}') ;")
            table_size_mb = float(cursor.fetchone()[0])/ (1024*1024)
            size_list.append(table_size_mb)

        tb_info = pd.DataFrame()
        tb_info["表名"] = tb_list
        tb_info["行数"] = row_count_list
        tb_info["大小（MB）"] = size_list
        tb_info.to_excel(self.out_file_path, index=False)

# 数据质检（是否有时间字段，是否有海洋数据标识码，是否矢量，矢量类型，最小线长度，最小面积）



# 测试
if __name__ == "__main__":
    test_instance = ReadDatabase()
    test_instance.input_file_path = "./测试数据/pg库表名.xlsx"
    #test_instance.out_file_path = "./测试数据/导出表格元数据.docx"
    #test_instance.out_file_path = "./测试数据/pg库表大小.xlsx"
    test_instance.out_file_path = "./测试数据/导出表格元数据/"

    test_instance.host = 'localhost'
    test_instance.database = 'pg_test'
    test_instance.user = 'postgres'
    test_instance.password = '314159'
    test_instance.port = '5432'

    #test_instance.get_tables_meta_data()
    test_instance.get_tables_meta_data_excel()
